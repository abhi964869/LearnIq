"""Document text extraction — PDF (PyMuPDF), DOCX (python-docx), TXT/MD."""
from __future__ import annotations

import io
import re

from .logger import get_logger

logger = get_logger(__name__)


class ExtractionError(Exception):
    """Raised when a document cannot be parsed."""


def extract_text(filename: str, data: bytes) -> tuple[str, dict]:
    """Extract plain text from an uploaded document.

    Args:
        filename: Original file name (extension decides the parser).
        data: Raw file bytes.

    Returns:
        (cleaned_text, metadata) where metadata includes pages/words.

    Raises:
        ExtractionError: For unsupported, empty, or corrupt files.
    """
    if not data:
        raise ExtractionError("The uploaded file is empty.")
    name = filename.lower()
    if name.endswith(".pdf"):
        text, pages = _extract_pdf(data)
    elif name.endswith(".docx"):
        text, pages = _extract_docx(data), None
    elif name.endswith((".txt", ".md")):
        text, pages = _decode_text(data), None
    else:
        raise ExtractionError(
            "Unsupported file type. Please upload a PDF, DOCX, TXT, or MD file."
        )
    text = clean_text(text)
    if len(text.split()) < 20:
        raise ExtractionError(
            "Could not extract readable text — the file may be scanned images "
            "or corrupt."
        )
    meta = {"words": len(text.split()), "chars": len(text)}
    if pages is not None:
        meta["pages"] = pages
    return text, meta


def clean_text(text: str) -> str:
    """Normalize whitespace, strip control characters, and de-hyphenate line breaks."""
    text = text.replace("\x00", "")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)          # de-hyphenate wraps
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(data: bytes) -> tuple[str, int]:
    """Extract text from PDF bytes using PyMuPDF."""
    import fitz  # PyMuPDF — imported lazily to keep cold starts fast for non-PDF calls

    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            pages = doc.page_count
            parts = [page.get_text("text") for page in doc]
    except Exception as exc:  # fitz raises generic RuntimeError on corrupt files
        logger.warning("PDF parse failure: %s", exc)
        raise ExtractionError("This PDF could not be opened — it may be corrupt "
                              "or password-protected.") from exc
    return "\n".join(parts), pages


def _extract_docx(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError("This DOCX file could not be opened.") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _decode_text(data: bytes) -> str:
    """Decode plain-text bytes with sensible fallbacks."""
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    raise ExtractionError("Could not decode this text file.")
